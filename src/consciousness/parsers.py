"""
Parsers for different data sources in Consciousness Capture component.
"""

import email
import json
import hashlib
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional
from email.utils import parsedate_to_datetime
import mimetypes


class EmailParser:
    """Parse email messages from various formats."""

    def __init__(self):
        self.supported_formats = ['.eml', '.mbox', '.msg']

    def parse_eml_file(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a single .eml email file.

        Args:
            file_path: Path to the .eml file

        Returns:
            Parsed email data dictionary
        """
        with open(file_path, 'rb') as f:
            msg = email.message_from_binary_file(f)

        return self._extract_message_data(msg)

    def parse_mbox_file(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Parse an mbox file containing multiple emails.

        Args:
            file_path: Path to the mbox file

        Returns:
            List of parsed email data dictionaries
        """
        import mailbox

        mbox = mailbox.mbox(file_path)
        messages = []

        for msg in mbox:
            try:
                messages.append(self._extract_message_data(msg))
            except Exception as e:
                # Log error but continue processing other messages
                print(f"Error parsing message: {e}")
                continue

        return messages

    def _extract_message_data(self, msg: email.message.Message) -> Dict[str, Any]:
        """
        Extract data from an email message object.

        Args:
            msg: Email message object

        Returns:
            Dictionary containing email data
        """
        # Extract headers
        headers = {
            'from': msg.get('From', ''),
            'to': msg.get('To', ''),
            'cc': msg.get('Cc', ''),
            'bcc': msg.get('Bcc', ''),
            'subject': msg.get('Subject', ''),
            'date': msg.get('Date', ''),
            'message_id': msg.get('Message-ID', ''),
            'in_reply_to': msg.get('In-Reply-To', ''),
            'references': msg.get('References', ''),
        }

        # Parse date
        timestamp = None
        if headers['date']:
            try:
                timestamp = parsedate_to_datetime(headers['date']).isoformat()
            except:
                timestamp = headers['date']

        # Extract body
        body_plain = ''
        body_html = ''
        attachments = []

        for part in msg.walk():
            content_type = part.get_content_type()
            content_disposition = str(part.get('Content-Disposition', ''))

            # Skip container parts
            if part.is_multipart():
                continue

            # Extract text content
            if content_type == 'text/plain' and 'attachment' not in content_disposition:
                try:
                    body_plain += part.get_payload(decode=True).decode('utf-8', errors='ignore')
                except:
                    pass

            elif content_type == 'text/html' and 'attachment' not in content_disposition:
                try:
                    body_html += part.get_payload(decode=True).decode('utf-8', errors='ignore')
                except:
                    pass

            # Handle attachments
            elif 'attachment' in content_disposition:
                filename = part.get_filename()
                if filename:
                    attachments.append({
                        'filename': filename,
                        'content_type': content_type,
                        'size': len(part.get_payload(decode=True)) if part.get_payload(decode=True) else 0
                    })

        # Create unique ID
        content_hash = hashlib.sha256(
            (headers['message_id'] + body_plain + str(timestamp)).encode()
        ).hexdigest()

        return {
            'id': content_hash[:16],
            'type': 'email',
            'timestamp': timestamp or datetime.now().isoformat(),
            'headers': headers,
            'content': {
                'plain': body_plain.strip(),
                'html': body_html.strip() if body_html else None,
            },
            'attachments': attachments,
            'metadata': {
                'char_count': len(body_plain),
                'has_attachments': len(attachments) > 0,
                'thread_id': headers.get('in_reply_to', ''),
            }
        }


class TextFileParser:
    """Parse various text file formats."""

    def __init__(self):
        self.supported_formats = [
            '.txt', '.md', '.markdown', '.rst', '.log',
            '.json', '.yaml', '.yml', '.toml', '.ini', '.cfg',
            '.py', '.js', '.java', '.cpp', '.c', '.h',
            '.html', '.xml', '.css', '.sql'
        ]

    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a text file.

        Args:
            file_path: Path to the text file

        Returns:
            Parsed file data dictionary
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Detect encoding
        encoding = self._detect_encoding(file_path)

        # Read content
        with open(file_path, 'r', encoding=encoding) as f:
            content = f.read()

        # Get file stats
        stats = path.stat()

        # Create unique ID
        content_hash = hashlib.sha256(content.encode()).hexdigest()

        return {
            'id': content_hash[:16],
            'type': 'text_file',
            'timestamp': datetime.fromtimestamp(stats.st_mtime).isoformat(),
            'content': {
                'text': content,
                'format': path.suffix.lower(),
            },
            'metadata': {
                'file_path': str(path.absolute()),
                'file_name': path.name,
                'size_bytes': stats.st_size,
                'created': datetime.fromtimestamp(stats.st_ctime).isoformat(),
                'modified': datetime.fromtimestamp(stats.st_mtime).isoformat(),
                'encoding': encoding,
                'line_count': content.count('\n') + 1,
                'word_count': len(content.split()),
                'char_count': len(content),
            }
        }

    def _detect_encoding(self, file_path: str) -> str:
        """
        Detect file encoding.

        Args:
            file_path: Path to file

        Returns:
            Detected encoding string
        """
        import chardet

        with open(file_path, 'rb') as f:
            raw_data = f.read(10000)  # Read first 10KB
            result = chardet.detect(raw_data)
            return result['encoding'] or 'utf-8'


class ChatParser:
    """Parse chat exports from various platforms."""

    def parse_whatsapp_export(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Parse WhatsApp chat export.

        Args:
            file_path: Path to WhatsApp export file

        Returns:
            List of parsed messages
        """
        import re

        messages = []
        pattern = r'(\d{1,2}/\d{1,2}/\d{2,4}, \d{1,2}:\d{2}(?:\s?[AP]M)?) - ([^:]+): (.+)'

        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        matches = re.findall(pattern, content, re.MULTILINE)

        for match in matches:
            timestamp_str, sender, message = match

            # Parse timestamp (handle different formats)
            try:
                # Try common formats
                for fmt in ['%m/%d/%y, %I:%M %p', '%d/%m/%y, %H:%M', '%m/%d/%Y, %I:%M %p']:
                    try:
                        timestamp = datetime.strptime(timestamp_str, fmt)
                        break
                    except:
                        continue
                else:
                    timestamp = datetime.now()
            except:
                timestamp = datetime.now()

            messages.append({
                'id': hashlib.sha256(f"{timestamp_str}{sender}{message}".encode()).hexdigest()[:16],
                'type': 'chat_message',
                'platform': 'whatsapp',
                'timestamp': timestamp.isoformat(),
                'sender': sender.strip(),
                'content': {
                    'text': message.strip()
                },
                'metadata': {
                    'char_count': len(message)
                }
            })

        return messages

    def parse_discord_export(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Parse Discord chat export (JSON format).

        Args:
            file_path: Path to Discord export file

        Returns:
            List of parsed messages
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        messages = []
        for msg in data.get('messages', []):
            messages.append({
                'id': msg.get('id', ''),
                'type': 'chat_message',
                'platform': 'discord',
                'timestamp': msg.get('timestamp', ''),
                'sender': msg.get('author', {}).get('name', ''),
                'content': {
                    'text': msg.get('content', '')
                },
                'metadata': {
                    'channel': data.get('channel', {}).get('name', ''),
                    'server': data.get('guild', {}).get('name', ''),
                    'attachments': len(msg.get('attachments', [])),
                    'reactions': len(msg.get('reactions', [])),
                }
            })

        return messages


class DocumentParser:
    """Parse document files (PDF, DOCX, etc.)."""

    def parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Parse PDF document.

        Args:
            file_path: Path to PDF file

        Returns:
            Parsed document data
        """
        try:
            import PyPDF2
        except ImportError:
            return {
                'id': 'pdf_error',
                'type': 'document',
                'format': 'pdf',
                'content': {'text': 'PyPDF2 not installed. Install with: pip install PyPDF2'},
                'metadata': {'file_path': file_path}
            }

        try:
            text_content = []
            metadata = {'pages': 0, 'file_path': file_path}

            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata['pages'] = len(pdf_reader.pages)

                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text:
                            text_content.append(f"[Page {page_num + 1}]\n{text}")
                    except Exception as e:
                        text_content.append(f"[Page {page_num + 1}] Error: {str(e)}")

                # Get document info
                if pdf_reader.metadata:
                    metadata['title'] = pdf_reader.metadata.get('/Title', '')
                    metadata['author'] = pdf_reader.metadata.get('/Author', '')
                    metadata['subject'] = pdf_reader.metadata.get('/Subject', '')

            combined_text = "\n\n".join(text_content)

            # Generate ID from content hash
            import hashlib
            content_hash = hashlib.sha256(combined_text.encode()).hexdigest()

            return {
                'id': content_hash[:16],
                'type': 'document',
                'format': 'pdf',
                'content': {
                    'text': combined_text,
                    'pages': text_content
                },
                'metadata': metadata
            }

        except Exception as e:
            return {
                'id': 'pdf_error',
                'type': 'document',
                'format': 'pdf',
                'content': {'text': f'Error parsing PDF: {str(e)}'},
                'metadata': {'file_path': file_path}
            }

    def parse_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Parse DOCX document.

        Args:
            file_path: Path to DOCX file

        Returns:
            Parsed document data
        """
        try:
            from docx import Document
        except ImportError:
            return {
                'id': 'docx_error',
                'type': 'document',
                'format': 'docx',
                'content': {'text': 'python-docx not installed. Install with: pip install python-docx'},
                'metadata': {'file_path': file_path}
            }

        try:
            doc = Document(file_path)

            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):
                        table_data.append(' | '.join(row_data))
                if table_data:
                    tables_text.append('\n'.join(table_data))

            # Combine all text
            all_text = '\n\n'.join(paragraphs)
            if tables_text:
                all_text += '\n\n[Tables]\n' + '\n\n'.join(tables_text)

            # Extract metadata
            metadata = {
                'file_path': file_path,
                'paragraphs': len(paragraphs),
                'tables': len(doc.tables)
            }

            # Try to get document properties
            try:
                core_props = doc.core_properties
                metadata['title'] = core_props.title or ''
                metadata['author'] = core_props.author or ''
                metadata['created'] = str(core_props.created) if core_props.created else ''
                metadata['modified'] = str(core_props.modified) if core_props.modified else ''
            except:
                pass

            # Generate ID from content hash
            import hashlib
            content_hash = hashlib.sha256(all_text.encode()).hexdigest()

            return {
                'id': content_hash[:16],
                'type': 'document',
                'format': 'docx',
                'content': {
                    'text': all_text,
                    'paragraphs': paragraphs,
                    'tables': tables_text
                },
                'metadata': metadata
            }

        except Exception as e:
            return {
                'id': 'docx_error',
                'type': 'document',
                'format': 'docx',
                'content': {'text': f'Error parsing DOCX: {str(e)}'},
                'metadata': {'file_path': file_path}
            }


class UniversalParser:
    """Universal parser that delegates to specific parsers."""

    def __init__(self):
        self.email_parser = EmailParser()
        self.text_parser = TextFileParser()
        self.chat_parser = ChatParser()
        self.document_parser = DocumentParser()

    def parse(self, file_path: str) -> Any:
        """
        Parse any supported file type.

        Args:
            file_path: Path to file

        Returns:
            Parsed data
        """
        path = Path(file_path)
        extension = path.suffix.lower()

        # Email formats
        if extension in ['.eml']:
            return self.email_parser.parse_eml_file(file_path)
        elif extension in ['.mbox']:
            return self.email_parser.parse_mbox_file(file_path)

        # Text formats
        elif extension in self.text_parser.supported_formats:
            return self.text_parser.parse_file(file_path)

        # Document formats
        elif extension == '.pdf':
            return self.document_parser.parse_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            return self.document_parser.parse_docx(file_path)

        # Chat exports (usually .txt with specific formatting)
        elif 'whatsapp' in path.name.lower():
            return self.chat_parser.parse_whatsapp_export(file_path)
        elif 'discord' in path.name.lower() and extension == '.json':
            return self.chat_parser.parse_discord_export(file_path)

        else:
            raise ValueError(f"Unsupported file format: {extension}")